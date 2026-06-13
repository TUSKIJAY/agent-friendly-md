"""Full Phase 0-7 smoke coverage for generated PDF and DOCX inputs."""
from __future__ import annotations

import importlib.util
import json
import subprocess
import sys
import tempfile
import unittest
import warnings
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(PROJECT_ROOT))
WRAPPER = PROJECT_ROOT / "scripts" / "run_python.py"

from lib import backend as bk  # noqa: E402
from lib import portability  # noqa: E402

warnings.filterwarnings(
    "ignore",
    category=DeprecationWarning,
    message="builtin type .* has no __module__ attribute",
)


def _deps_ok() -> bool:
    return all(importlib.util.find_spec(m) for m in ("docx", "fitz", "PIL", "playwright"))


def _doctor_full_ready() -> bool:
    spec = importlib.util.spec_from_file_location("doctor", PROJECT_ROOT / "tools" / "doctor.py")
    doctor = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(doctor)
    readiness = doctor.gather()["readiness"]
    return bool(readiness["core"] and readiness["extract"] and readiness["visual"])


def run(target: str, *args: str):
    return subprocess.run(
        [sys.executable, str(WRAPPER), target, *args],
        capture_output=True,
        text=True,
        encoding="utf-8",
        timeout=120,
    )


def _make_pdf(path: Path) -> None:
    import fitz  # noqa: PLC0415

    doc = fitz.open()
    page = doc.new_page(width=420, height=320)
    page.insert_text((48, 56), "Smoke PDF", fontsize=18)
    page.insert_text((48, 96), "This page exercises source-page visual evidence.", fontsize=11)
    page.insert_text((48, 126), "Formula: E = mc^2", fontsize=11)
    doc.save(path)
    doc.close()


def _make_docx(path: Path) -> None:
    from docx import Document  # noqa: PLC0415

    doc = Document()
    doc.add_heading("Smoke DOCX", level=1)
    doc.add_paragraph("This document exercises text source evidence.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "42"
    doc.save(path)


@unittest.skipUnless(bk.check_available()[0] and _deps_ok() and _doctor_full_ready(),
                     "full backend/extract/visual dependencies unavailable")
class TestFullSmokePipeline(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.root = Path(self.tmp.name)
        self.jobs = self.root / "jobs"

    def tearDown(self):
        self.tmp.cleanup()

    def _run_full_pipeline(self, source: Path, slug: str) -> tuple[Path, Path, dict]:
        r = run("tools/init_document_job.py", "--source", str(source),
                "--jobs-root", str(self.jobs), "--slug", slug, "--date", "20260613")
        self.assertEqual(r.returncode, 0, r.stdout + r.stderr)
        job = self.jobs / f"{slug}-20260613"
        package = self.root / f"{slug}-package"

        steps = [
            ("tools/run_validators.py", "--job", str(job)),
            ("tools/run_extraction.py", "--job", str(job)),
            ("tools/run_validators.py", "--job", str(job)),
            ("tools/build_ir.py", "--job", str(job)),
            ("tools/link_assets.py", "--job", str(job)),
            ("tools/run_validators.py", "--job", str(job)),
            ("tools/compile_structure.py", "--job", str(job)),
            ("tools/normalize_tables.py", "--job", str(job)),
            ("tools/normalize_formulas.py", "--job", str(job)),
            ("tools/semanticize_figures.py", "--job", str(job)),
            ("tools/run_validators.py", "--job", str(job)),
            ("tools/run_audits.py", "--job", str(job)),
            ("tools/run_validators.py", "--job", str(job)),
            ("tools/assemble_agent_md.py", "--job", str(job)),
            ("tools/run_validators.py", "--job", str(job)),
            ("tools/build_visual_review_packet.py", "--job", str(job)),
            ("tools/run_validators.py", "--job", str(job)),
            ("tools/package_output.py", "--job", str(job), "--out", str(package), "--force"),
        ]
        for target, *args in steps:
            r = run(target, *args)
            allowed = (0, 2) if target == "tools/run_validators.py" else (0,)
            self.assertIn(r.returncode, allowed, f"{target} failed:\nSTDOUT:\n{r.stdout}\nSTDERR:\n{r.stderr}")

        diagnostics = json.loads(
            (job / "review/visual_acceptance/render_diagnostics.json").read_text(encoding="utf-8")
        )
        return job, package, diagnostics

    def test_generated_pdf_phase0_to_package(self):
        source = self.root / "smoke.pdf"
        _make_pdf(source)
        job, package, diagnostics = self._run_full_pipeline(source, "smoke-pdf")
        comparisons = diagnostics["source_comparisons"]
        self.assertTrue(comparisons["sampled_pages"])
        self.assertTrue(comparisons["artifacts"])
        self.assertTrue((package / "smoke.md").is_file())
        self.assertTrue((package / "_agent-md/index.md").is_file())
        self.assertTrue((package / "_agent-md/index.json").is_file())
        self.assertEqual(portability.scan_path_leaks(package), [])
        qa = (package / "_agent-md/qa_report.md").read_text(encoding="utf-8")
        self.assertIn("overall:", qa)
        state = json.loads((job / "STATE.json").read_text(encoding="utf-8"))
        self.assertEqual(state["phase_status"]["7_package"], "done")

    def test_generated_docx_phase0_to_package(self):
        source = self.root / "smoke.docx"
        _make_docx(source)
        job, package, diagnostics = self._run_full_pipeline(source, "smoke-docx")
        comparisons = diagnostics["source_comparisons"]
        self.assertTrue(comparisons["sampled_source_evidence"])
        self.assertTrue(comparisons["artifacts"])
        self.assertTrue((package / "smoke.md").is_file())
        self.assertTrue((package / "_agent-md/index.md").is_file())
        self.assertTrue((package / "_agent-md/index.json").is_file())
        self.assertEqual(portability.scan_path_leaks(package), [])
        state = json.loads((job / "STATE.json").read_text(encoding="utf-8"))
        self.assertEqual(state["document_profile"]["risk_level"], "high")
        self.assertEqual(state["phase_status"]["7_package"], "done")


if __name__ == "__main__":
    unittest.main()
